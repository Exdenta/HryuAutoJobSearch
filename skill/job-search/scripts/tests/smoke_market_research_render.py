#!/usr/bin/env python3
"""Offline smoke test for market_research_render.py.

Covers three scenarios:

  1. Full fixture — every schema key populated. Verify:
     - file exists & > 10 KB
     - heading text for "Executive Summary", "References", "Table of Contents"
       appears somewhere in the document
     - at least one Heading 1 style is applied to a paragraph
     - >= 2 tables (skill_table + salary_home_table)
     - every citation int referenced from sections[*].citations appears as
       "[n]" somewhere in the docx text
     - every source URL appears somewhere in the docx (body or hyperlink rel)

  2. `manager_report = None` — renderer produces a >1 KB docx without raising.

  3. Fixture with a source-less section and missing optional keys → no crash.

Run:
    python3 scripts/tests/smoke_market_research_render.py
"""
from __future__ import annotations

import re
import sys
import tempfile
import zipfile
from pathlib import Path

HERE = Path(__file__).resolve().parent
SCRIPTS = HERE.parent
sys.path.insert(0, str(SCRIPTS))

from docx import Document  # noqa: E402

from market_research import ResearchRun  # noqa: E402
from market_research_render import render_research_docx  # noqa: E402


failures: list[str] = []


def check(cond: bool, label: str) -> None:
    if cond:
        print(f"  OK  {label}")
    else:
        print(f"  FAIL {label}")
        failures.append(label)


# ---------------------------------------------------------------------------
# Full fixture — every schema key populated.
# ---------------------------------------------------------------------------

_GOOD_MANAGER_FIXTURE: dict = {
    "cover": {
        "title": "Market Research — frontend engineer in Berlin",
        "subtitle": "12-month forward view with 24-month historical context",
        "prepared_for": "Candidate",
        "prepared_on": "2026-04-20",
        "word_count_estimate": 2500,
    },
    "executive_summary": [
        "Demand for frontend engineers in Berlin is steady with 1,240 open roles.",
        "Compensation has been flat year-over-year at mid-level.",
        "Playwright and RSC are the highest-leverage skill gaps to close.",
    ],
    "key_findings": [
        "Mid-level compensation clusters at 72k EUR (P50).",
        "Amsterdam pays ~10% more nominally but costs ~18% more to live in.",
        "Top 10 employers account for ~40% of postings.",
        "LLM tooling fluency is now table-stakes on senior loops.",
        "Hiring loops average 4 weeks from screen to offer.",
    ],
    "sections": [
        {
            "id": "demand",
            "heading": "Market demand and volume",
            "paragraphs": [
                "The Berlin frontend market has ~1,240 active openings "
                "with healthy distribution across levels.",
                "Top employers include N26, Zalando, and Delivery Hero.",
            ],
            "bullets": [
                "Mid-level roles dominate (~500 openings).",
                "Senior roles see the tightest competition.",
            ],
            "citations": [1, 2],
        },
        {
            "id": "history",
            "heading": "24-month historical context",
            "paragraphs": [
                "Demand dropped in Q1 2024 after broad tech layoffs but has "
                "recovered steadily through 2025.",
            ],
            "bullets": ["Q1-2024 low point: ~720 openings."],
            "citations": [3],
        },
        {
            "id": "skills_match",
            "heading": "Your skills vs the market",
            "paragraphs": [
                "Vue + TypeScript is a strong foundation; React exposure and "
                "Playwright would broaden reachable roles significantly.",
            ],
            "skill_table": [
                {"skill": "Vue 3", "relevance": "strong", "demand": 6,
                 "notes": "Solid base but niche vs React."},
                {"skill": "React", "relevance": "critical", "demand": 9,
                 "notes": "Dominant in Berlin JDs."},
                {"skill": "Playwright", "relevance": "strong", "demand": 7,
                 "notes": "Replacing Cypress in most stacks."},
            ],
            "gap_skills": ["React", "Playwright", "GraphQL"],
            "overrepresented_skills": ["jQuery"],
            "citations": [4, 5],
        },
        {
            "id": "salary",
            "heading": "Compensation",
            "paragraphs": [
                "Mid-level base salaries sit at 60-84k EUR (P25-P75). "
                "Amsterdam pays more nominally but CoL eats the gap.",
            ],
            "salary_home_table": [
                {"level": "Junior", "currency": "EUR",
                 "p25_local": 45000, "p50_local": 52000, "p75_local": 58000,
                 "p25_usd":   48000, "p50_usd":   55000, "p75_usd":   62000},
                {"level": "Mid", "currency": "EUR",
                 "p25_local": 60000, "p50_local": 72000, "p75_local": 84000,
                 "p25_usd":   64000, "p50_usd":   77000, "p75_usd":   90000},
                {"level": "Senior", "currency": "EUR",
                 "p25_local": 80000, "p50_local": 95000, "p75_local": 115000,
                 "p25_usd":   85000, "p50_usd":  101000, "p75_usd":  123000},
            ],
            "salary_neighbors_table": [
                {"market": "Amsterdam", "level": "Mid", "currency": "EUR",
                 "p25_local": 62000, "p50_local": 75000, "p75_local": 90000,
                 "p25_usd":   66000, "p50_usd":   80000, "p75_usd":   96000,
                 "cost_of_living_index_vs_home": 1.18},
            ],
            "citations": [6, 7],
        },
        {
            "id": "upskilling",
            "heading": "Recommended upskilling",
            "paragraphs": ["Prioritise Playwright (fastest ROI), then React."],
            "plan_bullets": [
                "Week 1-2: Playwright basics.",
                "Week 3-6: React fundamentals + hooks.",
                "Week 7-10: GraphQL integration projects.",
            ],
            "citations": [8],
        },
    ],
    "recommendations": [
        {"priority": "must", "text": "Learn Playwright in 4 weeks.",
         "rationale": "Appears on 70% of mid-level JDs.", "citations": [5]},
        {"priority": "should", "text": "Ship a public React side project.",
         "rationale": "Signals transferable framework literacy.",
         "citations": [4]},
        {"priority": "nice", "text": "Publish an RSC-focused blog post.",
         "rationale": "Visibility in the current trend cycle.", "citations": [2]},
    ],
    "risks": [
        "Code-gen tools may commoditise CRUD UI work.",
        "Berlin salary stagnation could persist into 2027.",
    ],
    "opportunities": [
        "LLM-adjacent frontend roles pay a 10-20% premium.",
        "EU-wide remote roles broaden the addressable market.",
    ],
    "sources": [
        {"n": 1, "title": "LinkedIn Jobs — Frontend Berlin",
         "url": "https://linkedin.com/jobs/berlin-frontend",
         "date": "2026-03-28", "snippet": "1240 active roles."},
        {"n": 2, "title": "Vercel blog — RSC adoption",
         "url": "https://vercel.com/blog/rsc-adoption-2026",
         "date": "2026-02-15", "snippet": "Framework convergence."},
        {"n": 3, "title": "Layoffs.fyi — Germany",
         "url": "https://layoffs.fyi/germany", "date": "2026-01-10",
         "snippet": "2024 layoff cohort."},
        {"n": 4, "title": "Stack Overflow Developer Survey 2025",
         "url": "https://survey.stackoverflow.co/2025",
         "date": "2025-07-01", "snippet": "Framework share."},
        {"n": 5, "title": "Playwright adoption report",
         "url": "https://playwright.dev/blog/2026-adoption",
         "date": "2026-02-28", "snippet": "Replacing Cypress."},
        {"n": 6, "title": "Levels.fyi — Berlin frontend",
         "url": "https://levels.fyi/t/software-engineer/locations/berlin",
         "date": "2026-03-15", "snippet": "P50 72k EUR."},
        {"n": 7, "title": "Levels.fyi — Amsterdam frontend",
         "url": "https://levels.fyi/t/software-engineer/locations/amsterdam",
         "date": "2026-03-15", "snippet": "P50 75k EUR."},
        {"n": 8, "title": "Playwright docs — Intro",
         "url": "https://playwright.dev/docs/intro", "date": "",
         "snippet": "Getting started."},
    ],
    "manager_confidence": "medium",
    "gaps_acknowledged": ["salary_neighbors limited to Amsterdam only"],
}


def _good_run() -> ResearchRun:
    return ResearchRun(
        status="ok",
        workers_ok=["demand", "history", "skills_match", "salary_home",
                    "salary_neighbors", "upskilling"],
        workers_failed=[],
        worker_results={},
        manager_report=_GOOD_MANAGER_FIXTURE,
        elapsed_ms=12345,
        started_at_iso="2026-04-20T10:00:00Z",
        finished_at_iso="2026-04-20T10:02:05Z",
        location_used="Berlin",
        resume_sha1="abcdef0123456789" + "0" * 24,
        prefs_sha1="1234abcd" + "0" * 32,
        model="opus",
        error=None,
    )


def _doc_text(path: Path) -> str:
    """Extract all text from a docx including body + hyperlink rels."""
    d = Document(str(path))
    parts: list[str] = []
    for para in d.paragraphs:
        parts.append(para.text)
    for table in d.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    parts.append(para.text)
    # Also grab hyperlink targets from relationships (URLs may live there, not
    # in run text, if only icon-style links were inserted).
    with zipfile.ZipFile(path, "r") as zf:
        for name in zf.namelist():
            if name.endswith(".rels") or name.endswith(".xml"):
                try:
                    parts.append(zf.read(name).decode("utf-8", "ignore"))
                except Exception:
                    pass
    return "\n".join(parts)


# ---------------------------------------------------------------------------
# Scenario 1: full fixture
# ---------------------------------------------------------------------------
print("1. render_research_docx — full fixture")

tmp = Path(tempfile.mkdtemp(prefix="mr_render_"))
out1 = tmp / "full.docx"
run1 = _good_run()
result_path = render_research_docx(run1, out1)
check(result_path == out1, f"returns out_path (got {result_path})")
check(out1.exists(), "file exists")
size1 = out1.stat().st_size if out1.exists() else 0
check(size1 > 10_000, f"file > 10 KB (got {size1})")

d1 = Document(str(out1))

# Heading 1 style present
h1_names = {"Heading 1"}
h1_paras = [p for p in d1.paragraphs if p.style and p.style.name in h1_names]
check(len(h1_paras) >= 1, f"at least 1 Heading 1 paragraph (got {len(h1_paras)})")

# Required heading text present
body_text_paragraphs = "\n".join(p.text for p in d1.paragraphs)
full_text = _doc_text(out1)
for needle in ("Executive Summary", "References", "Table of Contents"):
    check(needle in body_text_paragraphs,
          f"heading text {needle!r} appears in document body")

# Table count >= 2
check(len(d1.tables) >= 2,
      f"at least 2 tables present (got {len(d1.tables)})")

# Citation markers
expected_cits: set[int] = set()
for sec in _GOOD_MANAGER_FIXTURE["sections"]:
    for c in sec.get("citations") or []:
        expected_cits.add(c)
for rec in _GOOD_MANAGER_FIXTURE["recommendations"]:
    for c in rec.get("citations") or []:
        expected_cits.add(c)
for c in sorted(expected_cits):
    check(f"[{c}]" in full_text,
          f"citation marker [{c}] appears in document text")

# Source URLs present
for src in _GOOD_MANAGER_FIXTURE["sources"]:
    url = src["url"]
    check(url in full_text, f"source URL present: {url}")


# ---------------------------------------------------------------------------
# Scenario 2: manager_report is None
# ---------------------------------------------------------------------------
print("\n2. render_research_docx — catastrophic (manager_report=None)")

out2 = tmp / "none.docx"
run2 = ResearchRun(
    status="failed",
    workers_ok=[],
    workers_failed=[
        {"topic": "demand", "status": "failed", "error_head": "cli_missing"},
        {"topic": "history", "status": "overall_timeout", "error_head": "deadline"},
    ],
    worker_results={},
    manager_report=None,
    elapsed_ms=2000,
    started_at_iso="2026-04-20T10:00:00Z",
    finished_at_iso="2026-04-20T10:00:02Z",
    location_used="Berlin",
    resume_sha1="deadbeef" + "0" * 32,
    prefs_sha1="cafebabe" + "0" * 32,
    model="opus",
    error="all workers failed",
)

try:
    rp2 = render_research_docx(run2, out2)
    exc2 = None
except Exception as e:  # pragma: no cover — should never raise
    rp2 = None
    exc2 = e

check(exc2 is None, f"no exception raised on None manager_report (got {exc2!r})")
check(out2.exists(), "fallback docx file exists")
size2 = out2.stat().st_size if out2.exists() else 0
check(size2 > 1000, f"fallback file > 1 KB (got {size2})")

# Metadata appendix should mention workers_failed topics.
txt2 = _doc_text(out2)
check("Report generation failed" in txt2, "fallback document contains failure notice")
check("demand" in txt2, "fallback appendix mentions failed topic 'demand'")


# ---------------------------------------------------------------------------
# Scenario 2b: manager_report is None BUT workers succeeded — fallback body path
# ---------------------------------------------------------------------------
print("\n2b. render_research_docx — fallback from worker results (manager timeout)")

out2b = tmp / "fallback.docx"

# Minimal-but-realistic worker_results covering the most important topics.
# Each has a `sources` list so the fallback's global dedup + citation
# numbering is exercised.
_FALLBACK_WORKER_RESULTS: dict = {
    "demand": {
        "role_family": "Senior Frontend Engineer",
        "headline_summary": "Demand is stable in Berlin with roughly 420 open senior roles.",
        "total_open_postings_estimate": 420,
        "postings_by_level": {"junior": 60, "mid": 180, "senior": 140, "lead": 40},
        "top_employers": [
            {"name": "Zalando", "count": 25},
            {"name": "N26", "count": 14},
        ],
        "confidence": "medium",
        "sources": [
            {"title": "LinkedIn Berlin senior frontend roles",
             "url": "https://linkedin.com/jobs/berlin-senior-frontend",
             "date": "2026-04-15", "snippet": "420 open roles"},
            {"title": "Zalando careers",
             "url": "https://jobs.zalando.com/en/?q=frontend",
             "date": "2026-04-18", "snippet": "25 open frontend roles"},
        ],
    },
    "history": {
        "headline": "Demand rebounded in 2025 after a 2024 trough.",
        "timeline": [
            {"quarter": "2024-Q2", "demand_index": 40,
             "notable_event": "Klarna layoffs"},
            {"quarter": "2025-Q4", "demand_index": 85, "notable_event": ""},
        ],
        "layoff_events": [
            {"company": "Klarna", "date": "2024-05", "count": 700},
        ],
        "hiring_freezes": ["Delivery Hero 2024 Q3"],
        "confidence": "medium",
        "sources": [
            # Deliberately duplicates demand[0] — dedup should collapse.
            {"title": "LinkedIn Berlin senior frontend roles",
             "url": "https://linkedin.com/jobs/berlin-senior-frontend",
             "date": "2026-04-15", "snippet": "same URL, should dedup"},
            {"title": "Layoffs.fyi Germany 2024",
             "url": "https://layoffs.fyi/germany",
             "date": "2024-05", "snippet": ""},
        ],
    },
    "current_trends": {
        "headline_summary": "RSC + AI copilots dominate the 2026 JD language.",
        "hot_topics": [
            {"topic": "React Server Components", "why": "half of new listings mention RSC"},
        ],
        "fading_topics": [
            {"topic": "Redux", "why": "Zustand and Jotai are replacing it in JDs"},
        ],
        "buzzwords_in_jds": ["AI-assisted", "LLM", "RSC"],
        "confidence": "high",
        "sources": [
            {"title": "Vercel blog — 2026 state of React",
             "url": "https://vercel.com/blog/state-of-react-2026",
             "date": "2026-03", "snippet": ""},
        ],
    },
    "salary_home": {
        "headline": "Senior frontend in Berlin: €70-95k median.",
        "home_table": [
            {"level": "mid", "currency": "EUR",
             "p25_local": 55000, "p50_local": 65000, "p75_local": 75000,
             "p25_usd": 60000, "p50_usd": 70000, "p75_usd": 81000},
            {"level": "senior", "currency": "EUR",
             "p25_local": 70000, "p50_local": 82000, "p75_local": 95000,
             "p25_usd": 76000, "p50_usd": 89000, "p75_usd": 103000},
        ],
        "confidence": "medium",
        "sources": [
            {"title": "levels.fyi Berlin",
             "url": "https://levels.fyi/locations/berlin",
             "date": "2026-04", "snippet": ""},
        ],
    },
}

run2b = ResearchRun(
    status="partial",
    workers_ok=["demand", "history", "current_trends", "salary_home"],
    workers_failed=[],
    worker_results=_FALLBACK_WORKER_RESULTS,
    manager_report=None,
    elapsed_ms=840_000,
    started_at_iso="2026-04-20T10:00:00Z",
    finished_at_iso="2026-04-20T10:14:00Z",
    location_used="Berlin, Germany",
    resume_sha1="abcd" + "0" * 36,
    prefs_sha1="9999" + "0" * 36,
    model="opus",
    error="manager: timed out after 1500s",
)

try:
    rp2b = render_research_docx(run2b, out2b)
    exc2b = None
except Exception as e:
    rp2b = None
    exc2b = e

check(exc2b is None, f"no exception when manager None + workers present (got {exc2b!r})")
check(out2b.exists(), "fallback body docx file exists")
size2b = out2b.stat().st_size if out2b.exists() else 0
# A worker-body fallback with 4 topics + dedup references should be noticeably
# bigger than the 1 KB failure notice — expect at least 15 KB of real content.
check(size2b > 15_000,
      f"fallback-from-workers file > 15 KB (got {size2b})")

txt2b = _doc_text(out2b)
# Fallback-specific notice
check("Synthesis step incomplete" in txt2b,
      "fallback document contains 'Synthesis step incomplete' notice")
# Each worker's headline should surface
check("Demand is stable in Berlin" in txt2b,
      "fallback renders demand headline")
check("RSC + AI copilots" in txt2b,
      "fallback renders current_trends headline")
# Topic headings from _TOPIC_HEADINGS
check("Market demand & volume" in txt2b,
      "demand topic heading rendered")
check("Compensation — home market" in txt2b,
      "salary_home topic heading rendered")
# Sources deduped across workers: the duplicated LinkedIn URL must appear
# exactly once in the References section. We look at the body URL text.
d2b = Document(str(out2b))
body_text2b = "\n".join(p.text for p in d2b.paragraphs)
linkedin_occurrences = body_text2b.count("linkedin.com/jobs/berlin-senior-frontend")
check(linkedin_occurrences == 1,
      f"duplicated source URL deduped to a single reference (count={linkedin_occurrences})")
# Global citation renumbering — 6 raw sources across 4 workers, 1 duplicate
# URL → 5 unique references. Expect [1]..[5] and NOT [6].
check("[1]" in body_text2b and "[5]" in body_text2b,
      "global citation numbers [1]..[5] present in body text")
check("[6]" not in body_text2b,
      "no extra citation [6] after dedup (6 raw sources → 5 unique)")


# ---------------------------------------------------------------------------
# Scenario 3: sparse fixture — missing optional keys, empty/missing sub-sections
# ---------------------------------------------------------------------------
print("\n3. render_research_docx — sparse fixture (missing optional keys)")

sparse: dict = {
    "cover": {
        "title": "Sparse Report",
        "subtitle": "",
        "prepared_for": "Candidate",
        "prepared_on": "2026-04-20",
        "word_count_estimate": 100,
    },
    "executive_summary": ["Only one bullet of summary."],
    "key_findings": [],
    "sections": [
        # Section with only a heading + empty paragraphs; no bullets, no table.
        {"id": "demand", "heading": "Lonely section",
         "paragraphs": [], "citations": []},
        # Section whose citations refer to n=999 (invalid) — must be dropped.
        {"id": "history", "heading": "Citations gone bad",
         "paragraphs": ["Something happened."],
         "citations": [999, 1]},
        # Section with a skill_table but no content-field metadata otherwise.
        {"id": "skills_match", "heading": "Skills",
         "paragraphs": ["Skills stuff."],
         "skill_table": [
             {"skill": "X", "relevance": "strong", "demand": 5, "notes": "n"},
         ],
         "citations": []},
    ],
    "recommendations": [],   # empty → renderer prints "No recommendations"
    "risks": [],
    "opportunities": [],
    "sources": [
        {"n": 1, "title": "Only source", "url": "https://example.com/only",
         "date": "", "snippet": ""},
    ],
    "manager_confidence": "low",
    "gaps_acknowledged": [],
}

run3 = ResearchRun(
    status="partial",
    workers_ok=["demand"],
    workers_failed=[],
    worker_results={},
    manager_report=sparse,
    elapsed_ms=1000,
    started_at_iso="2026-04-20T10:00:00Z",
    finished_at_iso="2026-04-20T10:00:01Z",
    location_used="",
    resume_sha1="",
    prefs_sha1="",
    model="opus",
    error=None,
)

out3 = tmp / "sparse.docx"
try:
    rp3 = render_research_docx(run3, out3)
    exc3 = None
except Exception as e:
    rp3 = None
    exc3 = e
check(exc3 is None, f"sparse fixture does not raise (got {exc3!r})")
check(out3.exists() and out3.stat().st_size > 1000,
      f"sparse docx written (size={out3.stat().st_size if out3.exists() else 0})")

# Invalid citation 999 must NOT appear in doc text.
txt3 = _doc_text(out3)
# Look for the literal "[999]" token (not inside any URL or rel).
# Strip out .xml.rels content first to be safe — we only care about body text.
d3 = Document(str(out3))
body_text3 = "\n".join(p.text for p in d3.paragraphs)
for t in d3.tables:
    for row in t.rows:
        for cell in row.cells:
            body_text3 += "\n" + cell.text
check("[999]" not in body_text3,
      f"invalid citation [999] was filtered out (body_text search)")
check("[1]" in body_text3,
      "valid citation [1] present in body text")


# ---------------------------------------------------------------------------
# Summary
# ---------------------------------------------------------------------------
print()
if failures:
    print(f"FAIL {len(failures)} check(s):")
    for f in failures:
        print(f"   - {f}")
    sys.exit(1)
print("OK  All market_research_render smoke checks passed.")

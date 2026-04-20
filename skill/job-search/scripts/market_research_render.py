"""DOCX renderer for `market_research.ResearchRun`.

Produces a polished Word report. Follows `skills/docx/SKILL.md`: Arial 11pt
body, built-in Heading 1/2 (for TOC outlineLevel), DXA table widths,
ShadingType.CLEAR headers, no unicode-bullet glyphs. Never raises: on
`manager_report is None` falls back to a minimal cover + failure notice so
the bot can always attach a document.
"""
from __future__ import annotations

import logging
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.document import Document as _DocObj
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Emu
from docx.table import Table, _Cell

from market_research import ResearchRun

log = logging.getLogger(__name__)

# Page / table geometry (DXA; 1440 DXA == 1 inch)
_DXA_PER_INCH = 1440
_PAGE_W = 12240              # 8.5"
_PAGE_H = 15840              # 11"
_MARGIN = 1440               # 1"
_CONTENT_W = _PAGE_W - 2 * _MARGIN  # 9360 DXA = 6.5"

_BODY_FONT = "Arial"
_BODY_SIZE_PT = 11
_H1_SIZE_PT = 20
_H2_SIZE_PT = 15
_TITLE_SIZE_PT = 32
_SUBTITLE_SIZE_PT = 16
_HEADER_FILL = "E8ECEF"      # light gray shade for table headers
_LINK_COLOR = RGBColor(0x1A, 0x5B, 0xBE)


# XML-safety helpers

def _xml_safe(s: Any) -> str:
    """Strip NULs + invalid XML control chars; coerce non-strings via str()."""
    if s is None:
        return ""
    if not isinstance(s, str):
        try:
            s = str(s)
        except Exception:
            return ""
    out = []
    for ch in s:
        cp = ord(ch)
        if cp == 0x00:
            continue
        if cp < 0x20 and ch not in ("\t", "\n", "\r"):
            out.append(" ")
            continue
        if 0xD800 <= cp <= 0xDFFF:  # unpaired surrogate
            out.append(" ")
            continue
        out.append(ch)
    return "".join(out)


def _fmt_int(n: Any) -> str:
    """Format ints with thousands separators; fall through to _xml_safe otherwise."""
    try:
        if isinstance(n, bool):
            return str(n)
        if isinstance(n, int):
            return f"{n:,}"
        if isinstance(n, float):
            return f"{int(n):,}"
        return _xml_safe(n)
    except Exception:
        return _xml_safe(n)


# Document setup

def _force_font_east_asia(style: Any, font_name: str) -> None:
    """Patch XML so East-Asian fallback also uses `font_name`. Best-effort."""
    try:
        rPr = style.element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rFonts.set(qn(attr), font_name)
    except Exception:
        pass


def _configure_styles(doc: _DocObj) -> None:
    """Arial 11pt default, US Letter geometry, bold Heading 1/2 in near-black."""
    # Page size (US Letter) + 1" margins on every section.
    for section in doc.sections:
        section.page_width = Emu(_PAGE_W * 635)   # 1 DXA = 635 EMU
        section.page_height = Emu(_PAGE_H * 635)
        section.top_margin = Emu(_MARGIN * 635)
        section.bottom_margin = Emu(_MARGIN * 635)
        section.left_margin = Emu(_MARGIN * 635)
        section.right_margin = Emu(_MARGIN * 635)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = _BODY_FONT
    normal.font.size = Pt(_BODY_SIZE_PT)
    _force_font_east_asia(normal, _BODY_FONT)

    for (name, size) in (("Heading 1", _H1_SIZE_PT), ("Heading 2", _H2_SIZE_PT)):
        try:
            s = styles[name]
        except KeyError:
            continue
        s.font.name = _BODY_FONT
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
        _force_font_east_asia(s, _BODY_FONT)

    for name in ("List Bullet", "List Number"):
        try:
            s = styles[name]
            s.font.name = _BODY_FONT
            s.font.size = Pt(_BODY_SIZE_PT)
            _force_font_east_asia(s, _BODY_FONT)
        except KeyError:
            pass


# Paragraph / run helpers

def _add_para(doc: _DocObj, text: str, *, style: str | None = None,
              bold: bool = False, italic: bool = False,
              size_pt: float | None = None,
              align: int | None = None,
              color: RGBColor | None = None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(_xml_safe(text))
        run.bold = bold
        run.italic = italic
        if size_pt is not None:
            run.font.size = Pt(size_pt)
        if color is not None:
            run.font.color.rgb = color
        run.font.name = _BODY_FONT
    return p


def _add_heading(doc: _DocObj, text: str, level: int = 1):
    """Heading via built-in style so Word's TOC field picks it up."""
    return doc.add_heading(_xml_safe(text) or " ", level=level)


def _add_bullet(doc: _DocObj, text: str, *, style: str = "List Bullet"):
    """Bulleted paragraph; fallback '- ' prefix if style is missing."""
    text = _xml_safe(text)
    if not text:
        return None
    try:
        return doc.add_paragraph(text, style=style)
    except KeyError:
        return doc.add_paragraph("- " + text)


def _add_hyperlink(paragraph, url: str, text: str) -> None:
    """Insert a clickable hyperlink run (with its own rId relationship)."""
    url = _xml_safe(url)
    text = _xml_safe(text) or url
    if not url:
        paragraph.add_run(text)
        return
    try:
        r_id = paragraph.part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
    except Exception:
        paragraph.add_run(text)
        return

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), _BODY_FONT)
    rFonts.set(qn("w:hAnsi"), _BODY_FONT)
    rPr.append(rFonts)
    c = OxmlElement("w:color")
    c.set(qn("w:val"), "%02X%02X%02X" % (_LINK_COLOR[0], _LINK_COLOR[1], _LINK_COLOR[2]))
    rPr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    new_run.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


# Citations

def _citation_marker(citations: Iterable[Any], valid_ns: set[int]) -> str:
    """Render a list of citation ints as a trailing `[1][2][3]` string.
    Silently drop numbers not in `valid_ns`."""
    out = []
    seen: set[int] = set()
    for c in citations or ():
        if not isinstance(c, int) or isinstance(c, bool):
            continue
        if c in seen:
            continue
        if c not in valid_ns:
            continue
        seen.add(c)
        out.append(f"[{c}]")
    return " " + "".join(out) if out else ""


# Table helpers

def _apply_cell_shading(cell: _Cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")  # ShadingType.CLEAR
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def _set_cell_width_dxa(cell: _Cell, dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tcW = tc_pr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tc_pr.append(tcW)
    tcW.set(qn("w:type"), "dxa")
    tcW.set(qn("w:w"), str(dxa))


def _set_cell_margins_dxa(cell: _Cell, top: int = 80, bottom: int = 80,
                           left: int = 120, right: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom),
                      ("left", left), ("right", right)):
        e = OxmlElement(f"w:{side}")
        e.set(qn("w:w"), str(val))
        e.set(qn("w:type"), "dxa")
        mar.append(e)
    tc_pr.append(mar)


def _set_table_width_dxa(table: Table, dxa: int) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(dxa))


def _set_cell_text(cell: _Cell, text: str, *, bold: bool = False,
                   align: int | None = None) -> None:
    cell.text = ""  # clear the default empty paragraph content
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(_xml_safe(text))
    run.bold = bold
    run.font.name = _BODY_FONT
    run.font.size = Pt(_BODY_SIZE_PT)


def _add_table(doc: _DocObj, headers: list[str], rows: list[list[str]],
               col_widths_dxa: list[int],
               *, right_align_cols: set[int] | None = None) -> Table:
    """Create a bordered table. `col_widths_dxa` must sum to the table width."""
    right_align_cols = right_align_cols or set()
    total = sum(col_widths_dxa)
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    _set_table_width_dxa(table, total)

    # header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        c = hdr.cells[i]
        _set_cell_width_dxa(c, col_widths_dxa[i])
        _set_cell_margins_dxa(c)
        _apply_cell_shading(c, _HEADER_FILL)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_text(c, h, bold=True)

    # body rows
    for r, row in enumerate(rows, start=1):
        tr = table.rows[r]
        for i in range(len(headers)):
            c = tr.cells[i]
            _set_cell_width_dxa(c, col_widths_dxa[i])
            _set_cell_margins_dxa(c)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            val = row[i] if i < len(row) else ""
            align = (WD_ALIGN_PARAGRAPH.RIGHT
                     if i in right_align_cols else None)
            _set_cell_text(c, val, align=align)
    return table


# TOC field

def _add_toc_field(doc: _DocObj) -> None:
    """Insert Word's built-in TOC field. Word rebuilds it on open."""
    p = doc.add_paragraph()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), r'TOC \o "1-2" \h \z \u')
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), _BODY_FONT)
    rFonts.set(qn("w:hAnsi"), _BODY_FONT)
    rPr.append(rFonts)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = "Right-click and select 'Update Field' to populate the table of contents."
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    fld.append(r)
    p._p.append(fld)


# Section renderers

def _render_cover(doc: _DocObj, report: dict | None,
                  confidence: str | None) -> None:
    cover = (report or {}).get("cover") if report else None
    cover = cover if isinstance(cover, dict) else {}

    title = cover.get("title") or "Market Research Report"
    subtitle = cover.get("subtitle") or ""
    prepared_for = cover.get("prepared_for") or ""
    prepared_on = cover.get("prepared_on") or ""

    # Push title down a bit for visual balance.
    for _ in range(3):
        doc.add_paragraph()

    _add_para(doc, title, bold=True, size_pt=_TITLE_SIZE_PT,
              align=WD_ALIGN_PARAGRAPH.CENTER)
    if subtitle:
        _add_para(doc, subtitle, italic=True, size_pt=_SUBTITLE_SIZE_PT,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    if prepared_for:
        _add_para(doc, f"Prepared for: {prepared_for}",
                  align=WD_ALIGN_PARAGRAPH.CENTER)
    if prepared_on:
        _add_para(doc, f"Prepared on: {prepared_on}",
                  align=WD_ALIGN_PARAGRAPH.CENTER)
    if confidence:
        _add_para(doc, f"Confidence: {confidence}",
                  align=WD_ALIGN_PARAGRAPH.CENTER, italic=True,
                  size_pt=_BODY_SIZE_PT,
                  color=RGBColor(0x55, 0x55, 0x55))

    doc.add_page_break()


def _render_toc(doc: _DocObj) -> None:
    _add_heading(doc, "Table of Contents", level=1)
    _add_toc_field(doc)
    doc.add_page_break()


def _render_executive_summary(doc: _DocObj, items: list[str]) -> None:
    _add_heading(doc, "Executive Summary", level=1)
    for para in items or []:
        _add_para(doc, para)


def _render_key_findings(doc: _DocObj, items: list[str]) -> None:
    _add_heading(doc, "Key Findings", level=1)
    for b in items or []:
        _add_bullet(doc, b)


def _render_bullets(doc: _DocObj, items: list[str]) -> None:
    for b in items or []:
        _add_bullet(doc, b)


def _fmt_col_index(col: Any) -> str:
    if isinstance(col, (int, float)) and not isinstance(col, bool):
        return f"{float(col):.2f}"
    return _xml_safe(col or "")


def _render_skill_table(doc: _DocObj, rows: list[dict]) -> None:
    if not rows:
        return
    headers = ["Skill", "Relevance", "Demand (0-10)", "Notes"]
    widths = [2000, 1600, 1760, 4000]  # sums to 9360 = _CONTENT_W
    body = [
        [
            _xml_safe(r.get("skill") or ""),
            _xml_safe(r.get("relevance") or ""),
            _fmt_int(r.get("demand")) if isinstance(r.get("demand"), int)
                else _xml_safe(r.get("demand") or ""),
            _xml_safe(r.get("notes") or ""),
        ]
        for r in rows if isinstance(r, dict)
    ]
    _add_table(doc, headers, body, widths, right_align_cols={2})


def _render_salary_home_table(doc: _DocObj, rows: list[dict]) -> None:
    if not rows:
        return
    headers = ["Level", "Currency", "P25 local", "P50 local", "P75 local",
               "P25 USD", "P50 USD", "P75 USD"]
    widths = [1360, 1000, 1200, 1200, 1200, 1200, 1200, 1000]  # 9360
    body = [
        [
            _xml_safe(r.get("level") or ""),
            _xml_safe(r.get("currency") or ""),
            _fmt_int(r.get("p25_local")), _fmt_int(r.get("p50_local")),
            _fmt_int(r.get("p75_local")),
            _fmt_int(r.get("p25_usd")), _fmt_int(r.get("p50_usd")),
            _fmt_int(r.get("p75_usd")),
        ]
        for r in rows if isinstance(r, dict)
    ]
    _add_table(doc, headers, body, widths, right_align_cols={2, 3, 4, 5, 6, 7})


def _render_salary_neighbors_table(doc: _DocObj, rows: list[dict]) -> None:
    if not rows:
        return
    headers = ["Market", "Level", "Currency",
               "P25 local", "P50 local", "P75 local",
               "P25 USD", "P50 USD", "P75 USD", "CoL vs home"]
    widths = [1260, 820, 820, 900, 900, 900, 900, 900, 900, 1060]  # 9360
    body = [
        [
            _xml_safe(r.get("market") or r.get("market_name") or ""),
            _xml_safe(r.get("level") or ""),
            _xml_safe(r.get("currency") or ""),
            _fmt_int(r.get("p25_local")), _fmt_int(r.get("p50_local")),
            _fmt_int(r.get("p75_local")),
            _fmt_int(r.get("p25_usd")), _fmt_int(r.get("p50_usd")),
            _fmt_int(r.get("p75_usd")),
            _fmt_col_index(r.get("cost_of_living_index_vs_home")),
        ]
        for r in rows if isinstance(r, dict)
    ]
    _add_table(doc, headers, body, widths,
               right_align_cols={3, 4, 5, 6, 7, 8, 9})


def _render_section(doc: _DocObj, sec: dict, valid_source_ns: set[int]) -> None:
    if not isinstance(sec, dict):
        return
    heading = sec.get("heading") or "Section"
    _add_heading(doc, heading, level=1)

    sid = sec.get("id") or ""
    paragraphs = sec.get("paragraphs") or []
    citations = sec.get("citations") or []
    cit_str = _citation_marker(citations, valid_source_ns)

    # Render paragraphs; append citation marker to LAST non-empty paragraph.
    last_idx = -1
    for i, p in enumerate(paragraphs):
        if isinstance(p, str) and p.strip():
            last_idx = i
    for i, p in enumerate(paragraphs):
        text = _xml_safe(p or "")
        if i == last_idx and cit_str:
            text = text + cit_str
        _add_para(doc, text)

    # Bulleted variants — most sections use `bullets`; upskilling uses
    # `plan_bullets`; skills_match uses `gap_skills` + `overrepresented_skills`.
    bullets = sec.get("bullets") or []
    if bullets:
        _render_bullets(doc, bullets)

    plan_bullets = sec.get("plan_bullets") or []
    if plan_bullets:
        _render_bullets(doc, plan_bullets)

    # skills_match: special labels for gap / overrepresented lists
    gap_skills = sec.get("gap_skills") or []
    overrep = sec.get("overrepresented_skills") or []
    if gap_skills:
        _add_para(doc, "Gap skills:", bold=True)
        _render_bullets(doc, gap_skills)
    if overrep:
        _add_para(doc, "Overrepresented skills:", bold=True)
        _render_bullets(doc, overrep)

    # Tables
    if sid == "skills_match":
        _render_skill_table(doc, sec.get("skill_table") or [])
    if sid == "salary":
        _render_salary_home_table(doc, sec.get("salary_home_table") or [])
        _render_salary_neighbors_table(doc, sec.get("salary_neighbors_table") or [])

    # If no paragraphs at all but citation marker exists, add a small footer.
    if not paragraphs and cit_str:
        _add_para(doc, cit_str.strip(), italic=True, size_pt=9,
                  color=RGBColor(0x55, 0x55, 0x55))


def _render_recommendations(doc: _DocObj, recs: list[dict],
                            valid_source_ns: set[int]) -> None:
    _add_heading(doc, "Recommendations", level=1)
    if not recs:
        _add_para(doc, "No recommendations provided.", italic=True)
        return
    groups: dict[str, list[dict]] = {"must": [], "should": [], "nice": []}
    for r in recs:
        if not isinstance(r, dict):
            continue
        pr = r.get("priority")
        if pr in groups:
            groups[pr].append(r)

    labels = {"must": "Must", "should": "Should", "nice": "Nice-to-have"}
    for key in ("must", "should", "nice"):
        items = groups.get(key) or []
        if not items:
            continue
        _add_heading(doc, labels[key], level=2)
        for r in items:
            text = _xml_safe(r.get("text") or "")
            rationale = _xml_safe(r.get("rationale") or "")
            cit_str = _citation_marker(r.get("citations") or [], valid_source_ns)
            # Build bullet with bold label.
            try:
                p = doc.add_paragraph(style="List Bullet")
            except KeyError:
                p = doc.add_paragraph()
                p.add_run("- ")
            lbl = p.add_run(f"{labels[key]}: ")
            lbl.bold = True
            body = text
            if rationale:
                body = f"{text} — {rationale}" if text else rationale
            if cit_str:
                body = body + cit_str
            p.add_run(body)


def _render_risks_opportunities(doc: _DocObj, risks: list[str],
                                opportunities: list[str]) -> None:
    if risks:
        _add_heading(doc, "Risks", level=2)
        _render_bullets(doc, risks)
    if opportunities:
        _add_heading(doc, "Opportunities", level=2)
        _render_bullets(doc, opportunities)


def _render_references(doc: _DocObj, sources: list[dict]) -> None:
    _add_heading(doc, "References", level=1)
    if not sources:
        _add_para(doc, "No sources recorded.", italic=True)
        return
    for src in sources:
        if not isinstance(src, dict):
            continue
        n = src.get("n")
        title = _xml_safe(src.get("title") or "(untitled)")
        url = _xml_safe(src.get("url") or "")
        date = _xml_safe(src.get("date") or "")
        n_str = f"[{n}] " if isinstance(n, int) and not isinstance(n, bool) else ""
        p = doc.add_paragraph()
        run = p.add_run(f"{n_str}{title}.")
        run.font.name = _BODY_FONT
        run.font.size = Pt(_BODY_SIZE_PT)
        if date:
            p.add_run(f" {date}.")
        if url:
            p.add_run(" ")
            _add_hyperlink(p, url, url)


def _render_metadata_appendix(doc: _DocObj, run: ResearchRun) -> None:
    _add_heading(doc, "Appendix: Run Metadata", level=2)
    wf = run.workers_failed or []
    wf_summary = ", ".join(
        f"{_xml_safe(f.get('topic', '?'))}:{_xml_safe(f.get('status', '?'))}"
        for f in wf
    ) or "none"
    parts = [
        f"model: {_xml_safe(run.model or '?')}",
        f"elapsed_ms: {run.elapsed_ms}",
        f"started_at: {_xml_safe(run.started_at_iso or '?')}",
        f"finished_at: {_xml_safe(run.finished_at_iso or '?')}",
        f"workers_ok: {len(run.workers_ok or [])}",
        f"workers_failed: {wf_summary}",
        f"location_used: {_xml_safe(run.location_used or '?')}",
        f"resume_sha1: {_xml_safe((run.resume_sha1 or '')[:8])}",
        f"prefs_sha1: {_xml_safe((run.prefs_sha1 or '')[:8])}",
        f"status: {_xml_safe(run.status or '?')}",
    ]
    _add_para(doc, " | ".join(parts), italic=True, size_pt=9,
              color=RGBColor(0x55, 0x55, 0x55))


# ---------------------------------------------------------------------------
# Fallback rendering — used when the manager synthesis failed but one or more
# workers succeeded. We take each worker's raw JSON output and render it as
# its own section, with a globally-deduplicated + renumbered References
# section at the end. The result is less polished than the manager path
# (no executive summary, no cross-topic recommendations), but a 14-minute
# run that captured 10 topical reports isn't wasted.
# ---------------------------------------------------------------------------

_TOPIC_HEADINGS: dict[str, str] = {
    "demand":            "Market demand & volume",
    "history":           "24-month historical context",
    "current_trends":    "Current industry trends",
    "skills_match":      "Your skills vs. the market",
    "projections":       "12-18 month projections",
    "salary_home":       "Compensation — home market",
    "salary_neighbors":  "Compensation — neighboring markets",
    "companies":         "Company landscape",
    "hiring_bar":        "Interview & hiring bar",
    "upskilling":        "Recommended upskilling",
}

# Canonical order matches market_research.WORKERS so the document flows the
# same way as the manager-synthesized version.
_FALLBACK_TOPIC_ORDER: tuple[str, ...] = (
    "demand", "history", "current_trends", "skills_match", "projections",
    "salary_home", "salary_neighbors", "companies", "hiring_bar", "upskilling",
)


def _normalize_url_for_dedup(url: str) -> str:
    """Canonical form for per-run source dedup — lowercase scheme+host+path,
    drop query and fragment. Same rule the manager prompt uses."""
    if not isinstance(url, str):
        return ""
    u = url.strip()
    if not u:
        return ""
    # Lowercase the scheme and host-path separator boundary without pulling
    # in urllib — simple manual split is enough for dedup.
    try:
        # scheme://host/path?query#frag
        scheme_sep = u.find("://")
        if scheme_sep < 0:
            return u.lower().split("?", 1)[0].split("#", 1)[0].rstrip("/")
        scheme = u[:scheme_sep].lower()
        rest = u[scheme_sep + 3:]
        rest = rest.split("?", 1)[0].split("#", 1)[0]
        rest = rest.rstrip("/").lower()
        return f"{scheme}://{rest}"
    except Exception:
        return u.lower()


def _collect_fallback_sources(
    worker_results: dict[str, dict],
) -> tuple[list[dict], dict[tuple[str, int], int]]:
    """Aggregate and dedup sources across all worker outputs.

    Returns `(sources, idx_map)` where:
      - `sources` is a list of `{n, title, url, date, snippet}` dicts, numbered
        from 1, in first-seen order across _FALLBACK_TOPIC_ORDER.
      - `idx_map` maps `(topic, worker_source_index_0_based)` → global `n`
        so we can rewrite inline `(src: N)` / `source_idx: N` citations in
        each worker's payload.
    """
    sources: list[dict] = []
    url_to_n: dict[str, int] = {}
    idx_map: dict[tuple[str, int], int] = {}
    next_n = 1
    for topic in _FALLBACK_TOPIC_ORDER:
        wr = worker_results.get(topic)
        if not isinstance(wr, dict):
            continue
        raw_sources = wr.get("sources")
        if not isinstance(raw_sources, list):
            continue
        for i, src in enumerate(raw_sources):
            if not isinstance(src, dict):
                continue
            url = src.get("url") or ""
            key = _normalize_url_for_dedup(url)
            if not key:
                continue
            if key in url_to_n:
                idx_map[(topic, i)] = url_to_n[key]
                continue
            n = next_n
            next_n += 1
            url_to_n[key] = n
            idx_map[(topic, i)] = n
            sources.append({
                "n": n,
                "title": src.get("title") or "(untitled)",
                "url": url,
                "date": src.get("date") or "",
                "snippet": src.get("snippet") or "",
            })
    return sources, idx_map


def _render_demand_body(doc: _DocObj, wr: dict) -> None:
    total = wr.get("total_open_postings_estimate")
    role_family = wr.get("role_family")
    if isinstance(total, int) and total:
        prefix = f"Role family: {role_family}." if role_family else ""
        _add_para(doc, f"{prefix} Total open postings estimate: {_fmt_int(total)}.".strip())
    levels = wr.get("postings_by_level")
    if isinstance(levels, dict) and levels:
        level_line = ", ".join(
            f"{k}: {_fmt_int(levels.get(k))}"
            for k in ("junior", "mid", "senior", "lead")
            if levels.get(k) is not None
        )
        if level_line:
            _add_para(doc, f"By level — {level_line}.")
    top = wr.get("top_employers")
    if isinstance(top, list) and top:
        _add_heading(doc, "Top employers", level=2)
        rows = [
            [_xml_safe(e.get("name") or ""), _fmt_int(e.get("count"))]
            for e in top if isinstance(e, dict)
        ]
        if rows:
            _add_table(doc, ["Employer", "Open roles"], rows,
                        col_widths_dxa=[int(_CONTENT_W * 0.75), int(_CONTENT_W * 0.25)],
                        right_align_cols={1})


def _render_history_body(doc: _DocObj, wr: dict) -> None:
    timeline = wr.get("timeline")
    if isinstance(timeline, list) and timeline:
        _add_heading(doc, "Demand timeline", level=2)
        rows = [
            [
                _xml_safe(t.get("quarter") or ""),
                _fmt_int(t.get("demand_index")),
                _xml_safe(t.get("notable_event") or ""),
            ]
            for t in timeline if isinstance(t, dict)
        ]
        if rows:
            _add_table(
                doc, ["Quarter", "Demand idx (0-100)", "Notable event"], rows,
                col_widths_dxa=[
                    int(_CONTENT_W * 0.15),
                    int(_CONTENT_W * 0.18),
                    int(_CONTENT_W * 0.67),
                ],
                right_align_cols={1},
            )
    layoffs = wr.get("layoff_events")
    if isinstance(layoffs, list) and layoffs:
        _add_heading(doc, "Layoff events", level=2)
        for e in layoffs:
            if not isinstance(e, dict):
                continue
            line = " · ".join(
                part for part in (
                    _xml_safe(e.get("company") or ""),
                    _xml_safe(e.get("date") or ""),
                    f"{_fmt_int(e.get('count'))} affected" if e.get("count") else "",
                ) if part
            )
            if line:
                _add_bullet(doc, line)
    freezes = wr.get("hiring_freezes")
    if isinstance(freezes, list) and freezes:
        _add_heading(doc, "Hiring freezes", level=2)
        for f in freezes:
            if isinstance(f, str) and f.strip():
                _add_bullet(doc, _xml_safe(f))


def _render_current_trends_body(doc: _DocObj, wr: dict) -> None:
    hot = wr.get("hot_topics")
    if isinstance(hot, list) and hot:
        _add_heading(doc, "Hot topics", level=2)
        for t in hot:
            if not isinstance(t, dict):
                continue
            head = _xml_safe(t.get("topic") or "")
            why = _xml_safe(t.get("why") or "")
            if head:
                _add_bullet(doc, f"{head} — {why}" if why else head)
    fading = wr.get("fading_topics")
    if isinstance(fading, list) and fading:
        _add_heading(doc, "Fading topics", level=2)
        for t in fading:
            if not isinstance(t, dict):
                continue
            head = _xml_safe(t.get("topic") or "")
            why = _xml_safe(t.get("why") or "")
            if head:
                _add_bullet(doc, f"{head} — {why}" if why else head)
    buzz = wr.get("buzzwords_in_jds")
    if isinstance(buzz, list) and buzz:
        _add_heading(doc, "Buzzwords in JDs", level=2)
        _add_para(doc, ", ".join(_xml_safe(w) for w in buzz if isinstance(w, str)))


def _render_skills_match_body(doc: _DocObj, wr: dict) -> None:
    grades = wr.get("skill_grades")
    if isinstance(grades, list) and grades:
        _add_heading(doc, "Skill match", level=2)
        rows = [
            [
                _xml_safe(g.get("skill") or ""),
                _xml_safe(g.get("relevance") or ""),
                _fmt_col_index(g.get("market_demand_score")),
                _xml_safe(g.get("notes") or ""),
            ]
            for g in grades if isinstance(g, dict)
        ]
        if rows:
            _add_table(
                doc, ["Skill", "Relevance", "Demand 0-10", "Notes"], rows,
                col_widths_dxa=[
                    int(_CONTENT_W * 0.22),
                    int(_CONTENT_W * 0.15),
                    int(_CONTENT_W * 0.13),
                    int(_CONTENT_W * 0.50),
                ],
                right_align_cols={2},
            )
    gaps = wr.get("gap_skills")
    if isinstance(gaps, list) and gaps:
        _add_heading(doc, "Gap skills (worth adding)", level=2)
        for s in gaps:
            if isinstance(s, str) and s.strip():
                _add_bullet(doc, _xml_safe(s))
    over = wr.get("overrepresented_skills")
    if isinstance(over, list) and over:
        _add_heading(doc, "Overrepresented on your resume", level=2)
        for s in over:
            if isinstance(s, str) and s.strip():
                _add_bullet(doc, _xml_safe(s))


def _render_projections_body(doc: _DocObj, wr: dict) -> None:
    traj = wr.get("demand_trajectory")
    sal = wr.get("salary_trajectory")
    if isinstance(traj, str) and traj:
        _add_para(doc, f"Demand trajectory: {_xml_safe(traj)}.")
    if isinstance(sal, str) and sal:
        _add_para(doc, f"Salary trajectory: {_xml_safe(sal)}.")
    emerging = wr.get("emerging_roles")
    if isinstance(emerging, list) and emerging:
        _add_heading(doc, "Emerging roles", level=2)
        for r in emerging:
            if not isinstance(r, dict):
                continue
            title = _xml_safe(r.get("title") or "")
            desc = _xml_safe(r.get("description") or "")
            fit = r.get("fit_score")
            spec = " [speculative]" if r.get("speculative") else ""
            head = f"{title} (fit {fit}/5){spec}" if isinstance(fit, int) else f"{title}{spec}"
            if desc:
                _add_bullet(doc, f"{head} — {desc}")
            elif head:
                _add_bullet(doc, head)


def _render_salary_generic(doc: _DocObj, wr: dict, *, neighbors: bool) -> None:
    key = "neighbors_table" if neighbors else "home_table"
    table = wr.get(key)
    # Some prompts use different keys — try common variants.
    for alt in ("salary_table", "table"):
        if not isinstance(table, list):
            table = wr.get(alt) if isinstance(wr.get(alt), list) else table
    if isinstance(table, list) and table:
        if neighbors:
            headers = ["Market", "Level", "Cur", "P25", "P50", "P75",
                       "P25 USD", "P50 USD", "P75 USD", "CoL vs home"]
            widths = [int(_CONTENT_W * w) for w in
                      (0.15, 0.10, 0.06, 0.08, 0.08, 0.08, 0.09, 0.09, 0.09, 0.18)]
            right_align = {3, 4, 5, 6, 7, 8, 9}

            def row_of(r: dict) -> list[str]:
                return [
                    _xml_safe(r.get("market") or ""),
                    _xml_safe(r.get("level") or ""),
                    _xml_safe(r.get("currency") or ""),
                    _fmt_int(r.get("p25_local")),
                    _fmt_int(r.get("p50_local")),
                    _fmt_int(r.get("p75_local")),
                    _fmt_int(r.get("p25_usd")),
                    _fmt_int(r.get("p50_usd")),
                    _fmt_int(r.get("p75_usd")),
                    _xml_safe(r.get("cost_of_living_index_vs_home") or ""),
                ]
        else:
            headers = ["Level", "Cur", "P25", "P50", "P75",
                       "P25 USD", "P50 USD", "P75 USD"]
            widths = [int(_CONTENT_W * w) for w in
                      (0.16, 0.08, 0.10, 0.12, 0.12, 0.14, 0.14, 0.14)]
            right_align = {2, 3, 4, 5, 6, 7}

            def row_of(r: dict) -> list[str]:
                return [
                    _xml_safe(r.get("level") or ""),
                    _xml_safe(r.get("currency") or ""),
                    _fmt_int(r.get("p25_local")),
                    _fmt_int(r.get("p50_local")),
                    _fmt_int(r.get("p75_local")),
                    _fmt_int(r.get("p25_usd")),
                    _fmt_int(r.get("p50_usd")),
                    _fmt_int(r.get("p75_usd")),
                ]
        rows = [row_of(r) for r in table if isinstance(r, dict)]
        if rows:
            _add_table(doc, headers, rows, col_widths_dxa=widths,
                        right_align_cols=right_align)


def _render_companies_body(doc: _DocObj, wr: dict) -> None:
    for list_key in ("top_companies", "companies", "top_employers"):
        lst = wr.get(list_key)
        if isinstance(lst, list) and lst:
            _add_heading(doc, list_key.replace("_", " ").title(), level=2)
            for c in lst:
                if isinstance(c, dict):
                    name = _xml_safe(c.get("name") or "")
                    note = _xml_safe(c.get("note") or c.get("notes") or "")
                    count = c.get("count")
                    head = name
                    if isinstance(count, int):
                        head = f"{head} ({_fmt_int(count)} open roles)"
                    if note:
                        _add_bullet(doc, f"{head} — {note}" if head else note)
                    elif head:
                        _add_bullet(doc, head)
                elif isinstance(c, str) and c.strip():
                    _add_bullet(doc, _xml_safe(c))
            break


def _render_hiring_bar_body(doc: _DocObj, wr: dict) -> None:
    for arr_key in ("stages", "loop_stages", "common_requirements", "requirements"):
        arr = wr.get(arr_key)
        if isinstance(arr, list) and arr:
            _add_heading(doc, arr_key.replace("_", " ").title(), level=2)
            for s in arr:
                if isinstance(s, dict):
                    label = _xml_safe(s.get("name") or s.get("stage") or s.get("topic") or "")
                    desc = _xml_safe(s.get("description") or s.get("notes") or "")
                    if label:
                        _add_bullet(doc, f"{label} — {desc}" if desc else label)
                elif isinstance(s, str) and s.strip():
                    _add_bullet(doc, _xml_safe(s))


def _render_upskilling_body(doc: _DocObj, wr: dict) -> None:
    for arr_key in ("plan", "plan_bullets", "recommendations"):
        arr = wr.get(arr_key)
        if isinstance(arr, list) and arr:
            _add_heading(doc, arr_key.replace("_", " ").title(), level=2)
            for item in arr:
                if isinstance(item, dict):
                    head = _xml_safe(
                        item.get("skill") or item.get("name") or item.get("title") or ""
                    )
                    desc = _xml_safe(item.get("rationale") or item.get("description") or "")
                    if head or desc:
                        _add_bullet(doc, f"{head} — {desc}" if head and desc else (head or desc))
                elif isinstance(item, str) and item.strip():
                    _add_bullet(doc, _xml_safe(item))
            break


_TOPIC_RENDERERS = {
    "demand":            _render_demand_body,
    "history":           _render_history_body,
    "current_trends":    _render_current_trends_body,
    "skills_match":      _render_skills_match_body,
    "projections":       _render_projections_body,
    "salary_home":       lambda d, w: _render_salary_generic(d, w, neighbors=False),
    "salary_neighbors":  lambda d, w: _render_salary_generic(d, w, neighbors=True),
    "companies":         _render_companies_body,
    "hiring_bar":        _render_hiring_bar_body,
    "upskilling":        _render_upskilling_body,
}


def _worker_headline(wr: dict) -> str:
    """First-sentence summary for a worker output — try each common field."""
    for key in ("headline_summary", "headline", "narrative", "summary"):
        v = wr.get(key)
        if isinstance(v, str) and v.strip():
            return v.strip()
    return ""


def _render_worker_topic_section(
    doc: _DocObj, topic: str, wr: dict,
    idx_map: dict[tuple[str, int], int],
) -> None:
    heading = _TOPIC_HEADINGS.get(topic) or topic.replace("_", " ").title()
    _add_heading(doc, heading, level=1)
    headline = _worker_headline(wr)
    if headline:
        _add_para(doc, _xml_safe(headline))
    # Topic-specific structured content
    renderer = _TOPIC_RENDERERS.get(topic)
    if renderer is not None:
        try:
            renderer(doc, wr)
        except Exception:
            log.exception("market_research_render: fallback renderer for %s crashed", topic)
    # Confidence marker at the end, if present
    conf = wr.get("confidence")
    if isinstance(conf, str) and conf:
        _add_para(doc, f"Topic confidence: {_xml_safe(conf)}.",
                  italic=True, size_pt=9, color=RGBColor(0x55, 0x55, 0x55))
    # Citation footer: which globally-renumbered sources this worker used.
    worker_srcs = wr.get("sources") or []
    if isinstance(worker_srcs, list) and worker_srcs:
        ns: list[int] = []
        for i in range(len(worker_srcs)):
            n = idx_map.get((topic, i))
            if isinstance(n, int) and n not in ns:
                ns.append(n)
        if ns:
            marker = " ".join(f"[{n}]" for n in ns)
            _add_para(
                doc, f"Sources: {marker}",
                italic=True, size_pt=9, color=RGBColor(0x55, 0x55, 0x55),
            )


def _render_fallback_body(doc: _DocObj, run: ResearchRun) -> None:
    """Render a usable report when the manager failed but workers succeeded.

    Called from `render_research_docx` as a better alternative to the tiny
    'Report generation failed' page when `worker_results` is non-empty.
    """
    worker_results = run.worker_results or {}
    sources, idx_map = _collect_fallback_sources(worker_results)

    _render_cover(doc, None, confidence=None)

    _add_heading(doc, "Synthesis step incomplete", level=1)
    _add_para(
        doc,
        "The synthesizer pass did not complete in time, so this report is "
        "assembled from the raw outputs of the worker agents. Each section "
        "below is one agent's narrow-topic report, in the canonical order.",
        italic=True,
    )
    if run.error:
        _add_para(doc, f"Synthesizer error: {_xml_safe(run.error)}",
                  italic=True, size_pt=9, color=RGBColor(0x55, 0x55, 0x55))

    for topic in _FALLBACK_TOPIC_ORDER:
        wr = worker_results.get(topic)
        if isinstance(wr, dict):
            _render_worker_topic_section(doc, topic, wr, idx_map)

    _render_references(doc, sources)
    _render_metadata_appendix(doc, run)


# Public entry point

def render_research_docx(run: ResearchRun, out_path: Path) -> Path:
    """Render a ResearchRun to a polished .docx at `out_path`. Returns out_path."""
    out_path = Path(out_path)
    doc = Document()
    try:
        _configure_styles(doc)
    except Exception:
        log.exception("market_research_render: style configuration failed (continuing)")

    report = run.manager_report if isinstance(run.manager_report, dict) else None

    if report is None:
        # Manager synthesis missing. If we still have worker output, salvage
        # it into a structured — if less polished — report via the fallback
        # renderer. Only fall through to the minimal failure notice when there
        # is genuinely nothing to render.
        worker_results = run.worker_results if isinstance(run.worker_results, dict) else {}
        has_workers = any(
            isinstance(v, dict) and v for v in worker_results.values()
        )
        if has_workers:
            try:
                _render_fallback_body(doc, run)
            except Exception:
                log.exception("market_research_render: fallback body crashed")
                # On a fallback crash, still produce *something* the user can open.
                _add_heading(doc, "Report assembly failed", level=1)
                _add_para(
                    doc,
                    "The fallback renderer crashed while assembling worker outputs. "
                    "The raw worker JSON is preserved on disk for debugging.",
                )
                if run.error:
                    _add_para(doc, f"Error: {_xml_safe(run.error)}", italic=True)
                _render_metadata_appendix(doc, run)
        else:
            # Catastrophic path — no manager AND no workers.
            _render_cover(doc, None, confidence=None)
            _add_heading(doc, "Report generation failed", level=1)
            _add_para(
                doc,
                "The market research run did not produce a synthesizer report "
                "and no worker agents completed successfully. See the appendix "
                "below for details of which workers failed.",
            )
            if run.error:
                _add_para(doc, f"Error: {_xml_safe(run.error)}", italic=True)
            _render_metadata_appendix(doc, run)
        try:
            doc.save(str(out_path))
        except Exception:
            log.exception("market_research_render: failed to save fallback docx")
            raise
        return out_path

    # Build valid source `n` set for citation filtering.
    sources = report.get("sources") or []
    valid_source_ns: set[int] = set()
    for s in sources:
        if isinstance(s, dict):
            n = s.get("n")
            if isinstance(n, int) and not isinstance(n, bool):
                valid_source_ns.add(n)

    confidence = report.get("manager_confidence")
    if not isinstance(confidence, str):
        confidence = None

    _render_cover(doc, report, confidence=confidence)
    _render_toc(doc)
    _render_executive_summary(doc, report.get("executive_summary") or [])
    _render_key_findings(doc, report.get("key_findings") or [])
    for sec in report.get("sections") or []:
        _render_section(doc, sec, valid_source_ns)
    _render_recommendations(doc, report.get("recommendations") or [],
                             valid_source_ns)
    _render_risks_opportunities(
        doc, report.get("risks") or [], report.get("opportunities") or [],
    )
    _render_references(doc, sources)
    _render_metadata_appendix(doc, run)

    try:
        doc.save(str(out_path))
    except Exception:
        log.exception("market_research_render: save failed")
        raise
    return out_path
